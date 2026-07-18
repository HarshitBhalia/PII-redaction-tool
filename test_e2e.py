"""Comprehensive E2E test with ALL PII types including tricky cases."""
import requests
import json
import os

BASE = "http://localhost:5000"

# Create a thorough test document covering ALL PII types
from docx import Document

doc = Document()
doc.add_heading("Comprehensive PII Test Document", 0)

doc.add_heading("Section 1 — Personal Information", level=1)
doc.add_paragraph("Full Name: Rashi Patil")
doc.add_paragraph("Alias: Dr. Sanjay Kumar Mehta (Director)")
doc.add_paragraph("Mr. Amit Kumar Singh, Mrs. Priya Sharma, and Rohan Dey are attending.")

doc.add_heading("Section 2 — Contact Details", level=1)
doc.add_paragraph("Email: rashhi.patil@gmail.com")
doc.add_paragraph("Work Email: rohan.dey@techsolutions.co.in")
doc.add_paragraph("Mobile: +91 9876543210")
doc.add_paragraph("Mobile 2: 9123456789")
doc.add_paragraph("Landline: 022-28001234")
doc.add_paragraph("Office: 080-45053237")
doc.add_paragraph("Toll Free: 1800-123-4567")

doc.add_heading("Section 3 — Company Names", level=1)
doc.add_paragraph("Company: TechSolutions India Pvt Ltd")
doc.add_paragraph("Parent Company: Reliance Industries Limited")
doc.add_paragraph("Banker: HDFC Bank Ltd")
doc.add_paragraph("Auditor: Deloitte Haskins & Sells LLP")
doc.add_paragraph("Partner Firm: Kumar & Associates")

doc.add_heading("Section 4 — Physical Addresses", level=1)
doc.add_paragraph("Registered Office: 123 MG Road, Bengaluru, Karnataka 560001")
doc.add_paragraph("Corporate Office: Plot No. 45, Sector 12, Noida, Uttar Pradesh - 201301")
doc.add_paragraph("Branch: Flat 3B, Sunshine Apartments, Andheri West, Mumbai 400058")
doc.add_paragraph("Residence: House No. 12, Gandhi Nagar, New Delhi - 110031")
doc.add_paragraph("Factory: Door No. 5/2, 3rd Cross, Jayanagar, Bangalore 560041")
doc.add_paragraph("Warehouse: C-14, Defence Colony, New Delhi - 110024")
doc.add_paragraph("Office: B-204, Silver Oak Society, Thane, Maharashtra 400601")
doc.add_paragraph("Mailing Address: P.O. Box 1234, Connaught Place, New Delhi")

doc.add_heading("Section 5 — Government IDs", level=1)
doc.add_paragraph("PAN Card: ABCDE1234F")
doc.add_paragraph("PAN of Director: ZZZZZ9999Z")
doc.add_paragraph("Aadhaar Number: 1234 5678 9012")
doc.add_paragraph("Aadhaar (no spaces): 987654321098")
doc.add_paragraph("Passport Number: A1234567")
doc.add_paragraph("Driving License: MH12-20150012345")

doc.add_heading("Section 6 — Financial Information", level=1)
doc.add_paragraph("Credit Card: 4111-1111-1111-1111")
doc.add_paragraph("Card 2: 5500 0000 0000 0004")
doc.add_paragraph("SSN (US): 123-45-6789")
doc.add_paragraph("IFSC Code: HDFC0001234")

doc.add_heading("Section 7 — Digital Information", level=1)
doc.add_paragraph("Server IP: 192.168.1.100")
doc.add_paragraph("Public IP: 203.45.67.89")
doc.add_paragraph("Internal IP: 10.0.0.1")
doc.add_paragraph("Website: https://www.rashipatil.com")

doc.add_heading("Section 8 — Dates of Birth", level=1)
doc.add_paragraph("Date of Birth: 15/08/1990")
doc.add_paragraph("DOB (ISO): 1985-03-22")
doc.add_paragraph("Born on: March 22, 1985")
doc.add_paragraph("Birthday: 01 January 1975")

doc.add_paragraph("\nThis document is for testing the PII Redaction Tool.")
doc.save("comprehensive_pii_test.docx")
print("[OK] Created comprehensive_pii_test.docx")

# Upload
print("\n[1] Uploading...")
with open("comprehensive_pii_test.docx", "rb") as f:
    resp = requests.post(f"{BASE}/api/upload", files={"file": ("comprehensive_pii_test.docx", f)})
data = resp.json()
file_id = data['file_id']
print(f"    File ID: {file_id}")

# Redact
print("\n[2] Redacting...")
resp = requests.post(f"{BASE}/api/redact", json={
    "file_id": file_id,
    "filename": "comprehensive_pii_test.docx",
    "selected_types": [],
    "clear_mapping": True
})
data = resp.json()
print(f"    Status: {'OK' if data.get('success') else 'FAILED: ' + data.get('error','?')}")
stats = data.get('stats', {})
print(f"    Total PII found: {stats.get('total_findings', 0)}")
print(f"    Entity breakdown:")
for et, cnt in sorted(stats.get('entity_counts', {}).items()):
    print(f"      {et:20s}: {cnt}")
output_filename = data.get('output_filename')

# Download & Show
print(f"\n[3] Downloading {output_filename}...")
resp = requests.get(f"{BASE}/api/download/{output_filename}")
out_path = f"outputs/{output_filename}"
with open(out_path, 'wb') as f:
    f.write(resp.content)

print("\n[4] REDACTED CONTENT:")
print("=" * 60)
doc2 = Document(out_path)
for para in doc2.paragraphs:
    if para.text.strip():
        print(f"  {para.text}")

print("\n[DONE]")
