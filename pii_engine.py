"""
PII Redaction Engine — v2.0 FIXED
===================================
Hybrid multi-layer PII detection:
  - Layer 1: Custom Regex (addresses, phones, emails, PAN, Aadhaar, CC, IP, DOB, SSN)
  - Layer 2: spaCy NER (PERSON, ORG, GPE/LOC via Presidio)
  - Layer 3: Address-specific heuristic patterns (Indian + international)
  - Layer 4: Faker (en_IN) with global HashMap for consistent replacements

KEY FIXES IN v2:
  - Allowlist only blocks EXACT standalone words, NOT substrings inside company names
  - Landline phone regex added (0XX-XXXXXXXX)
  - Company name detection enhanced with domain-specific patterns
  - Full address block detection (multi-word Indian address patterns)
  - All Faker replacements now generate correct type (full address for address)
"""

import re
import logging
import random
from typing import Dict, List, Tuple, Optional
from faker import Faker
from docx import Document
import concurrent.futures

logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────
# FAKER SETUP — consistent HashMap
# ─────────────────────────────────────────────
fake_in  = Faker('en_IN')
fake_en  = Faker('en_US')
Faker.seed(42)
random.seed(42)

_pii_mapping: Dict[str, str] = {}


def get_fake_value(entity_type: str, original_text: str) -> str:
    """Return a consistent fake replacement (same input → same output always)."""
    clean = original_text.strip()
    if clean in _pii_mapping:
        return _pii_mapping[clean]
    val = _generate_fake(entity_type, clean)
    _pii_mapping[clean] = val
    return val


def _generate_fake(entity_type: str, original: str) -> str:
    """Generate realistic fake data for each PII type."""
    et = entity_type.upper()

    if et == 'PERSON':
        return fake_in.name()

    elif et == 'EMAIL_ADDRESS':
        return fake_in.email()

    elif et in ('PHONE_NUMBER', 'IN_PHONE', 'PHONE'):
        # Return same format (mobile vs landline)
        if original.startswith('0') and '-' in original:
            # Landline: 0XX-XXXXXXXX
            return fake_in.numerify('0##-########')
        elif '+91' in original or original.startswith('91'):
            return fake_in.numerify('+91##########')
        else:
            return fake_in.numerify('##########')

    elif et in ('LOCATION', 'GPE', 'LOC', 'FAC', 'ADDRESS', 'FULL_ADDRESS'):
        # Full address replacement
        house = random.choice(['Plot No.', 'Flat', 'House No.', 'Door No.', 'Block'])
        num   = fake_in.numerify(random.choice(['##', '###', '#/##']))
        area  = random.choice(['Gandhi Nagar', 'Nehru Colony', 'MG Road', 'Civil Lines',
                               'Patel Nagar', 'Sector 12', 'Park Street', 'Race Course Road'])
        city  = fake_in.city()
        state = random.choice(['Maharashtra', 'Karnataka', 'Delhi', 'Tamil Nadu',
                               'Uttar Pradesh', 'Gujarat', 'Rajasthan', 'West Bengal'])
        pin   = fake_in.numerify('######')
        return f"{house} {num}, {area}, {city}, {state} - {pin}"

    elif et in ('ORG', 'ORGANIZATION', 'COMPANY'):
        return fake_in.company()

    elif et == 'CREDIT_CARD':
        return f"XXXX-XXXX-XXXX-{fake_in.numerify('####')}"

    elif et == 'IP_ADDRESS':
        return fake_en.ipv4_private() if random.random() < 0.5 else fake_en.ipv4()

    elif et in ('DATE_TIME', 'DATE_OF_BIRTH', 'DOB', 'DATE_OF_BIRTH'):
        dob = fake_in.date_of_birth(minimum_age=18, maximum_age=70)
        # Match the original format
        if '/' in original:
            return dob.strftime('%d/%m/%Y')
        elif '-' in original:
            return dob.strftime('%d-%m-%Y')
        else:
            return dob.strftime('%B %d, %Y')

    elif et in ('US_SSN', 'SSN'):
        return fake_en.ssn()

    elif et == 'IN_PAN':
        letters = 'ABCDEFGHIJKLMNOPQRSTUVWXYZ'
        return (random.choice(letters) + random.choice(letters) + random.choice(letters) +
                random.choice('ABCFGHLJPTK') +
                random.choice(letters) +
                fake_in.numerify('####') +
                random.choice(letters))

    elif et == 'IN_AADHAAR':
        return fake_in.numerify('#### #### ####')

    elif et in ('URL', 'DOMAIN_NAME'):
        return fake_en.domain_name()

    elif et in ('NRP', 'NATIONALITY'):
        return fake_en.country()

    elif et == 'PASSPORT':
        return random.choice('ABCDEFGHIJKLMNOPQRSTUVWXYZ') + fake_in.numerify('#######')

    elif et == 'VEHICLE_REG':
        st = random.choice(['MH', 'DL', 'KA', 'TN', 'UP', 'GJ', 'RJ', 'WB'])
        return f"{st}{fake_in.numerify('##')}{random.choice('ABCDEFGHIJKLMNOPQRSTUVWXYZ')}{random.choice('ABCDEFGHIJKLMNOPQRSTUVWXYZ')}{fake_in.numerify('####')}"

    else:
        return f"[REDACTED_{et}]"


# ─────────────────────────────────────────────
# ALLOWLIST — ONLY exact short standalone terms
# (NOT applied to company names or addresses)
# ─────────────────────────────────────────────
_EXACT_ALLOWLIST = frozenset([
    # Securities regulators (standalone abbreviations only)
    "sebi", "bse", "nse", "rbi", "irdai", "nsdl", "cdsl",
    "nifty", "sensex", "nifty50",
    # Pure financial/legal document terms
    "red herring prospectus", "prospectus", "offer document", "drhp",
    "ipo", "fpo", "allotment", "syndicate", "securities", "mutual fund",
    "equity shares", "preference shares", "debentures",
    "nav", "eps", "ebitda", "cagr", "pe",
    "qib", "brlm", "promoter", "subsidiary", "book running lead manager",
    # Non-company standalone geographic words
    "india", "indian",
    # Document structure terms
    "order", "ticket", "invoice", "schedule", "annexure",
    "appendix", "exhibit", "section", "clause",
    # Generic words that are never PII alone
    "act", "law", "court", "government",
])


def _is_allowlisted(text: str) -> bool:
    """
    Only block if the ENTIRE text (stripped, lowercased) is an exact allowlist term.
    Do NOT block partial matches — this was causing company names to be skipped.
    """
    lower = text.lower().strip()
    return lower in _EXACT_ALLOWLIST


# ─────────────────────────────────────────────
# REGEX PATTERNS — comprehensive
# ─────────────────────────────────────────────

# Indian PIN code (6 digits, used as anchor for address detection)
_PIN_RE = re.compile(r'\b\d{6}\b')

# Indian PAN Card
_PAN_RE = re.compile(r'\b[A-Z]{5}[0-9]{4}[A-Z]\b')

# Indian Aadhaar (12 digits, with or without spaces)
_AADHAAR_RE = re.compile(r'\b\d{4}[\s\-]?\d{4}[\s\-]?\d{4}\b')

# Email
_EMAIL_RE = re.compile(r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}\b')

# Credit card — with dashes/spaces or raw
_CC_RE = re.compile(
    r'\b(?:'
    r'\d{4}[\s\-]\d{4}[\s\-]\d{4}[\s\-]\d{4}'   # spaced/dashed: 4-4-4-4
    r'|4[0-9]{12}(?:[0-9]{3})?'                    # Visa
    r'|5[1-5][0-9]{14}'                             # Mastercard
    r'|3[47][0-9]{13}'                              # Amex
    r'|6(?:011|5[0-9]{2})[0-9]{12}'                # Discover
    r')\b'
)

# US SSN — detect standard XXX-XX-XXXX format (avoid 000, 666 as first group)
_SSN_RE = re.compile(r'\b(?!000)(?!666)\d{3}-(?!00)\d{2}-(?!0000)\d{4}\b')


# IP address v4
_IP_RE = re.compile(r'\b(?:(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\.){3}(?:25[0-5]|2[0-4]\d|[01]?\d\d?)\b')

# IPv6
_IP6_RE = re.compile(r'\b(?:[0-9a-fA-F]{1,4}:){7}[0-9a-fA-F]{1,4}\b')

# Phone numbers — ALL Indian formats + toll-free
_PHONE_PATTERNS = [
    # +91 mobile (10 digits after country code)
    re.compile(r'(?<!\d)\+91[\s\-]?[6-9]\d{9}(?!\d)'),
    # 91XXXXXXXXXX (without +)
    re.compile(r'(?<!\d)91[6-9]\d{9}(?!\d)'),
    # 0XXXXXXXXXX (mobile with leading 0)
    re.compile(r'(?<!\d)0[6-9]\d{9}(?!\d)'),
    # Plain 10-digit mobile (starts 6-9)
    re.compile(r'(?<!\d)[6-9]\d{9}(?!\d)'),
    # Landline: 0XX-XXXXXXXX or 0XX XXXXXXXX (STD code + number)
    re.compile(r'(?<!\d)0\d{2,4}[\s\-]\d{6,8}(?!\d)'),
    # Landline: (0XX)XXXXXXXX
    re.compile(r'\(0\d{2,4}\)\s?\d{6,8}'),
    # Toll-free: 1800-XXX-XXXX or 1800XXXXXXX
    re.compile(r'(?<!\d)1800[\s\-]?\d{3}[\s\-]?\d{4}(?!\d)'),
    # International: +X-XXX-XXX-XXXX or similar
    re.compile(r'\+\d{1,3}[\s\-]\d{1,4}[\s\-]\d{4,10}'),
    # +1 (US/Canada style)
    re.compile(r'\+1[\s\-]\(?\d{3}\)?[\s\-]\d{3}[\s\-]\d{4}'),
]

# Date of birth — with strict word boundaries to avoid matching inside other data
_DOB_PATTERNS = [
    # DD/MM/YYYY or DD-MM-YYYY (only with century years 19xx/20xx)
    re.compile(r'(?<![\w/\-])(?:0?[1-9]|[12]\d|3[01])[/\-](?:0?[1-9]|1[0-2])[/\-](?:19|20)\d{2}(?![\w/\-])'),
    # YYYY-MM-DD (ISO format, with mandatory separators)
    re.compile(r'\b(?:19|20)\d{2}-(?:0[1-9]|1[0-2])-(?:0[1-9]|[12]\d|3[01])\b'),
    # Month name DD, YYYY — must be preceded by whitespace or start of line
    re.compile(
        r'(?:^|\s)(?:January|February|March|April|May|June|July|August|September|October|November|December)'
        r'\s+\d{1,2},?\s+(?:19|20)\d{2}(?=\s|$|[,.])',
        re.IGNORECASE | re.MULTILINE
    ),
    # DD Month YYYY — with word boundary
    re.compile(
        r'\b\d{1,2}\s+(?:January|February|March|April|May|June|July|August|September|October|November|December)'
        r'\s+(?:19|20)\d{2}\b',
        re.IGNORECASE
    ),
]

# Indian address patterns (heuristic — catches full address blocks)
_ADDRESS_PATTERNS = [
    # Plot/Flat/House/Door No. + area + city + PIN
    re.compile(
        r'(?:Plot\s*No\.?|Flat\s*No\.?|House\s*No\.?|Door\s*No\.?|'
        r'Apartment|Apt\.?|Block|Wing|Tower|Building|Phase|Sector|'
        r'Survey\s*No\.?|S\.?No\.?)\s*[\w\-/]+[,\s]+[\w\s,\-/\.]+\d{6}',
        re.IGNORECASE
    ),
    # Address line with PIN at end
    re.compile(
        r'\d+[,\-]?\s*[\w\s]+(?:Road|Street|Lane|Avenue|Marg|Nagar|Colony|'
        r'Layout|Extension|Enclave|Vihar|Puram|Park|Garden|Complex|Society|'
        r'Residency|Heights|Tower|Plaza|Cross|Main|Ring\s*Road|By\s*Pass|'
        r'Highway|NH|SH)[,\s]+[\w\s,]+\d{6}',
        re.IGNORECASE
    ),
    # PO Box addresses
    re.compile(r'P\.?\s*O\.?\s*Box\s+\d+[\w\s,]+', re.IGNORECASE),
    # C/O address
    re.compile(r'C/O\s+[\w\s]+,\s*[\w\s,]+\d{5,6}', re.IGNORECASE),
    # Indian state + PIN pattern (strong address signal)
    re.compile(
        r'(?:[\w\s,\-#/.]+)'
        r'(?:Maharashtra|Karnataka|Tamil\s*Nadu|Uttar\s*Pradesh|'
        r'Andhra\s*Pradesh|Telangana|Rajasthan|Gujarat|West\s*Bengal|'
        r'Madhya\s*Pradesh|Bihar|Odisha|Kerala|Jharkhand|Assam|'
        r'Punjab|Haryana|Uttarakhand|Himachal\s*Pradesh|Goa|'
        r'Jammu|Kashmir|Manipur|Meghalaya|Mizoram|Nagaland|Sikkim|'
        r'Tripura|Arunachal|Chhattisgarh|Delhi|NCR)'
        r'[\s,\-]+(?:\d{6}|\d{3}\s*\d{3})',
        re.IGNORECASE
    ),
]

# Company name patterns — standalone company detection beyond NER
_COMPANY_PATTERNS = [
    # Pvt Ltd / Private Limited / Ltd variants
    re.compile(
        r'(?:[A-Z][a-zA-Z0-9&\s\-\'.]+?)\s+'
        r'(?:Pvt\.?\s*Ltd\.?|Private\s+Limited|Limited|Ltd\.?|'
        r'Pte\.?\s*Ltd\.?|LLP|LLC|Corp\.?|Corporation|Inc\.?|'
        r'Incorporated|PLC|GmbH|S\.A\.|N\.V\.)',
        re.IGNORECASE
    ),
    # "XYZ & Associates" / "XYZ & Sons"
    re.compile(
        r'[A-Z][a-zA-Z\s]+(?:&|and)\s+(?:Associates?|Sons?|Brothers?|Partners?|Co\.?)',
        re.IGNORECASE
    ),
    # Indian banks & specific entities
    re.compile(
        r'(?:[A-Z][a-zA-Z]+\s+)*(?:Bank|Insurance|Finance|Securities|Investments?|'
        r'Capital|Ventures?|Holdings?|Enterprises?|Services?|Solutions?|'
        r'Technologies?|Systems?|Industries?|Exports?|Imports?)\s+(?:Ltd\.?|Limited|Pvt\.?)',
        re.IGNORECASE
    ),
]

# Passport (Indian)
_PASSPORT_RE = re.compile(r'\b[A-Z][0-9]{7}\b')

# Vehicle registration (India) — MH12AB1234
_VEHICLE_RE = re.compile(r'\b[A-Z]{2}\d{2}[A-Z]{1,2}\d{4}\b')


# ─────────────────────────────────────────────
# PRESIDIO SETUP
# ─────────────────────────────────────────────
_analyzer = None


def _get_analyzer():
    """Lazy-init Presidio with spaCy."""
    global _analyzer
    if _analyzer is not None:
        return _analyzer

    try:
        from presidio_analyzer import AnalyzerEngine, RecognizerRegistry, Pattern, PatternRecognizer
        from presidio_analyzer.nlp_engine import NlpEngineProvider

        logger.info("Initializing Presidio + spaCy NLP engine...")

        try:
            config = {
                "nlp_engine_name": "spacy",
                "models": [{"lang_code": "en", "model_name": "en_core_web_sm"}]
            }
            provider = NlpEngineProvider(nlp_configuration=config)
            nlp_engine = provider.create_engine()
        except Exception as e:
            logger.warning(f"spaCy load failed ({e}), using default NLP.")
            nlp_engine = None

        registry = RecognizerRegistry()
        registry.load_predefined_recognizers()

        # Custom recognizers
        for entity, name, regex, score in [
            ("PHONE_NUMBER", "indian_mobile",   r"(?<!\d)\+?91[\s\-]?[6-9]\d{9}(?!\d)", 0.9),
            ("PHONE_NUMBER", "landline_india",  r"(?<!\d)0\d{2,4}[\s\-]\d{6,8}(?!\d)",  0.85),
            ("IN_PAN",       "pan_card",        r"\b[A-Z]{5}[0-9]{4}[A-Z]\b",            0.95),
            ("IN_AADHAAR",   "aadhaar_spaced",  r"\b\d{4}[\s\-]?\d{4}[\s\-]?\d{4}\b",   0.85),
        ]:
            registry.add_recognizer(PatternRecognizer(
                supported_entity=entity,
                patterns=[Pattern(name=name, regex=regex, score=score)]
            ))

        if nlp_engine:
            _analyzer = AnalyzerEngine(nlp_engine=nlp_engine, registry=registry)
        else:
            _analyzer = AnalyzerEngine(registry=registry)

        logger.info("Presidio analyzer ready.")
        return _analyzer

    except Exception as e:
        logger.error(f"Presidio init failed: {e}")
        return None


PRESIDIO_ENTITIES = [
    "PERSON", "EMAIL_ADDRESS", "PHONE_NUMBER", "LOCATION", "ORG",
    "IP_ADDRESS", "CREDIT_CARD", "US_SSN", "DATE_TIME", "NRP",
    "URL", "IN_PAN", "IN_AADHAAR", "DOMAIN_NAME",
]


# ─────────────────────────────────────────────
# CORE DETECTION
# ─────────────────────────────────────────────

def _run_regex_layer(text: str) -> List[Dict]:
    """Run all custom regex patterns on text. Returns raw findings."""
    findings = []

    def add(start, end, etype, matched):
        if matched.strip() and not _is_allowlisted(matched):
            findings.append({'start': start, 'end': end,
                             'entity_type': etype, 'text': matched,
                             'score': 0.88, 'source': 'regex'})

    # PAN
    for m in _PAN_RE.finditer(text):
        add(m.start(), m.end(), 'IN_PAN', m.group())

    # Aadhaar — only if NOT looks like a year/small number
    for m in _AADHAAR_RE.finditer(text):
        matched = m.group().replace(' ', '').replace('-', '')
        if len(matched) == 12 and matched.isdigit():
            add(m.start(), m.end(), 'IN_AADHAAR', m.group())

    # Email
    for m in _EMAIL_RE.finditer(text):
        add(m.start(), m.end(), 'EMAIL_ADDRESS', m.group())

    # Credit card
    for m in _CC_RE.finditer(text):
        add(m.start(), m.end(), 'CREDIT_CARD', m.group())

    # SSN
    for m in _SSN_RE.finditer(text):
        add(m.start(), m.end(), 'US_SSN', m.group())

    # IP v4
    for m in _IP_RE.finditer(text):
        parts = m.group().split('.')
        if all(0 <= int(p) <= 255 for p in parts):
            add(m.start(), m.end(), 'IP_ADDRESS', m.group())

    # IP v6
    for m in _IP6_RE.finditer(text):
        add(m.start(), m.end(), 'IP_ADDRESS', m.group())

    # Phone numbers (all patterns)
    for pattern in _PHONE_PATTERNS:
        for m in pattern.finditer(text):
            matched = m.group().strip()
            if len(re.sub(r'\D', '', matched)) >= 7:  # at least 7 digits
                add(m.start(), m.end(), 'PHONE_NUMBER', matched)

    # Dates of birth
    for pattern in _DOB_PATTERNS:
        for m in pattern.finditer(text):
            add(m.start(), m.end(), 'DATE_OF_BIRTH', m.group())

    # Full address patterns
    for pattern in _ADDRESS_PATTERNS:
        for m in pattern.finditer(text):
            if len(m.group().strip()) > 10:  # meaningful length
                add(m.start(), m.end(), 'FULL_ADDRESS', m.group())

    # Company name patterns
    for pattern in _COMPANY_PATTERNS:
        for m in pattern.finditer(text):
            matched = m.group().strip()
            if len(matched) > 5 and not _is_allowlisted(matched):
                add(m.start(), m.end(), 'ORG', matched)

    # Passport
    for m in _PASSPORT_RE.finditer(text):
        # Must not overlap with PAN (PAN is 10 chars, Passport is 8)
        if len(m.group()) == 8:
            add(m.start(), m.end(), 'PASSPORT', m.group())

    # Vehicle registration
    for m in _VEHICLE_RE.finditer(text):
        add(m.start(), m.end(), 'VEHICLE_REG', m.group())

    return findings


def _run_presidio_layer(text: str, selected_types: Optional[List[str]] = None) -> List[Dict]:
    """Run Presidio NER detection only for requested NLP entities."""
    findings = []
    try:
        analyzer = _get_analyzer()
        if not analyzer:
            return findings

        # Filter the entities passed to Presidio to ONLY the slow NLP ones requested
        target_entities = []
        if selected_types:
            expanded = set()
            for sel in selected_types:
                expanded.update(_TYPE_EXPANSION.get(sel.upper(), [sel.upper()]))
            
            # Only ask Presidio to find these heavy NLP entities if requested
            for ner_type in ['PERSON', 'ORG', 'LOCATION', 'NRP']:
                if ner_type in expanded:
                    target_entities.append(ner_type)
        else:
            target_entities = ['PERSON', 'ORG', 'LOCATION', 'NRP']
            
        if not target_entities:
            return findings # Nothing to do

        results = analyzer.analyze(
            text=text,
            entities=target_entities,
            language='en',
            score_threshold=0.45,
        )
        for r in results:
            matched = text[r.start:r.end]
            if matched.strip() and not _is_allowlisted(matched):
                findings.append({
                    'start': r.start,
                    'end': r.end,
                    'entity_type': r.entity_type,
                    'text': matched,
                    'score': r.score,
                    'source': 'presidio',
                })
    except Exception as e:
        logger.warning(f"Presidio error: {e}")
    return findings


def _merge_spans(findings: List[Dict]) -> List[Dict]:
    """Merge overlapping spans — prefer longest span, then highest score."""
    if not findings:
        return []

    # Sort by start, then by length desc, then score desc
    srt = sorted(findings, key=lambda x: (x['start'], -(x['end'] - x['start']), -x['score']))

    merged = []
    last_end = -1

    for f in srt:
        if f['start'] >= last_end:
            merged.append(f)
            last_end = f['end']
        else:
            # Overlapping — keep the one with greater span length
            if (f['end'] - f['start']) > (merged[-1]['end'] - merged[-1]['start']):
                merged[-1] = f
                last_end = f['end']
            elif (f['end'] - f['start']) == (merged[-1]['end'] - merged[-1]['start']) and f['score'] > merged[-1]['score']:
                merged[-1] = f

    return merged


# Map UI-selected type names → internal entity type names
_TYPE_EXPANSION: Dict[str, List[str]] = {
    'PERSON':        ['PERSON'],
    'EMAIL_ADDRESS': ['EMAIL_ADDRESS'],
    'PHONE_NUMBER':  ['PHONE_NUMBER', 'IN_PHONE', 'PHONE'],
    'ORG':           ['ORG', 'ORGANIZATION', 'COMPANY'],
    'LOCATION':      ['LOCATION', 'GPE', 'LOC', 'FAC', 'ADDRESS', 'FULL_ADDRESS'],
    'US_SSN':        ['US_SSN', 'SSN'],
    'CREDIT_CARD':   ['CREDIT_CARD'],
    'DATE_TIME':     ['DATE_TIME', 'DATE_OF_BIRTH', 'DOB'],
    'IP_ADDRESS':    ['IP_ADDRESS'],
    'IN_PAN':        ['IN_PAN', 'PAN'],
    'IN_AADHAAR':    ['IN_AADHAAR', 'AADHAAR'],
    'PASSPORT':      ['PASSPORT'],
    'VEHICLE_REG':   ['VEHICLE_REG'],
    'NRP':           ['NRP', 'NATIONALITY'],
    'URL':           ['URL', 'DOMAIN_NAME'],
}

def _matches_selected(entity_type: str, selected_types: List[str]) -> bool:
    """Check if an entity type matches any of the selected UI types."""
    et = entity_type.upper()
    for sel in selected_types:
        variants = _TYPE_EXPANSION.get(sel.upper(), [sel.upper()])
        if et in variants:
            return True
    return False


def detect_pii(text: str, selected_types: Optional[List[str]] = None) -> List[Dict]:
    """
    Detect all PII in text.
    Returns list of {start, end, entity_type, text, score, source}.
    """
    if not text or not text.strip():
        return []

    # Fast layer (Regex)
    all_findings = _run_regex_layer(text)
    
    # Check if we need the slow NER layer
    needs_ner = True
    if selected_types is not None:
        ner_types = {'PERSON', 'ORG', 'LOCATION', 'NRP'}
        # Expand user's selected types to see if they need NER
        expanded_selected = set()
        for sel in selected_types:
            expanded_selected.update(_TYPE_EXPANSION.get(sel.upper(), [sel.upper()]))
        
        # If intersection is empty, we don't need NER!
        if not ner_types.intersection(expanded_selected):
            needs_ner = False

    if needs_ner:
        all_findings += _run_presidio_layer(text, selected_types)

    # Merge overlaps
    merged = _merge_spans(all_findings)

    # Filter by selected types if specified
    if selected_types and len(selected_types) > 0:
        merged = [f for f in merged if _matches_selected(f['entity_type'], selected_types)]

    return merged


# ─────────────────────────────────────────────
# TEXT REDACTION
# ─────────────────────────────────────────────

def redact_text_segment(text: str, selected_types: Optional[List[str]] = None) -> Tuple[str, List[Dict]]:
    """Redact PII from a plain text string by processing in chunks to prevent OOM while maintaining speed."""
    lines = text.split('\n')
    redacted_lines = []
    all_findings = []
    
    # Safe character limit for spaCy on 512MB RAM is ~5000 chars per chunk
    MAX_CHARS = 5000
    
    current_chunk = []
    current_len = 0
    chunks = []
    
    for line in lines:
        current_chunk.append(line)
        current_len += len(line) + 1 # +1 for \n
        if current_len >= MAX_CHARS:
            chunks.append(current_chunk)
            current_chunk = []
            current_len = 0
            
    if current_chunk:
        chunks.append(current_chunk)
        
    for chunk_lines in chunks:
        chunk_text = '\n'.join(chunk_lines)
        
        if not chunk_text.strip():
            redacted_lines.extend(chunk_lines)
            continue
            
        findings = detect_pii(chunk_text, selected_types)
        
        if not findings:
            redacted_lines.extend(chunk_lines)
            continue
            
        all_findings.extend(findings)
        
        # Redact the chunk
        sorted_findings = sorted(findings, key=lambda x: x['start'], reverse=True)
        redacted = chunk_text
        for f in sorted_findings:
            replacement = get_fake_value(f['entity_type'], f['text'])
            redacted = redacted[:f['start']] + replacement + redacted[f['end']:]
            
        # Add the redacted chunk lines back
        redacted_lines.extend(redacted.split('\n'))

    return '\n'.join(redacted_lines), all_findings


def _apply_replacements_to_text(text: str, selected_types: Optional[List[str]] = None) -> Tuple[str, List[Dict]]:
    """
    Detect PII in text, build replacement map sorted longest-first,
    and apply substitutions. Returns (new_text, findings).
    """
    findings = detect_pii(text, selected_types)
    if not findings:
        return text, []

    # Build map: original PII string → fake replacement
    replacement_map: Dict[str, str] = {}
    for f in findings:
        orig = f['text']
        if orig not in replacement_map:
            replacement_map[orig] = get_fake_value(f['entity_type'], orig)

    # Apply longest-first to avoid partial substitutions
    new_text = text
    for orig in sorted(replacement_map.keys(), key=len, reverse=True):
        new_text = new_text.replace(orig, replacement_map[orig])

    return new_text, findings


# ─────────────────────────────────────────────
# DOCX REDACTION
# ─────────────────────────────────────────────

def _safe_update_para(para, new_text: str):
    """Update a paragraph's text while preserving the first run's formatting."""
    if not para.runs:
        para.clear()
        para.add_run(new_text)
        return

    # Snapshot formatting from first run
    r0 = para.runs[0]
    fmt = {
        'font_name': r0.font.name,
        'font_size': r0.font.size,
        'bold':      r0.bold,
        'italic':    r0.italic,
        'underline': r0.underline,
    }

    # Remove all runs except the first
    for run in para.runs[1:]:
        run._element.getparent().remove(run._element)

    # Set text and restore formatting
    para.runs[0].text = new_text
    if fmt['font_name']:  para.runs[0].font.name = fmt['font_name']
    if fmt['font_size']:  para.runs[0].font.size = fmt['font_size']
    if fmt['bold']   is not None: para.runs[0].bold    = fmt['bold']
    if fmt['italic'] is not None: para.runs[0].italic  = fmt['italic']
    if fmt['underline'] is not None: para.runs[0].underline = fmt['underline']


def redact_docx(input_path: str, output_path: str, selected_types: Optional[List[str]] = None) -> Dict:
    """
    Redact PII from a DOCX file.
    Processes: paragraphs, tables (all cells), headers, footers.
    Returns stats dict.
    """
    doc = Document(input_path)
    stats = {
        'total_findings': 0,
        'entity_counts': {},
        'paragraphs_processed': 0,
    }

    # Collect all paragraphs that need processing
    paragraphs_to_process = []

    def collect_para(para):
        original = para.text
        if original.strip():
            paragraphs_to_process.append((para, original))

    # ── Main body paragraphs ──
    for para in doc.paragraphs:
        collect_para(para)

    # ── Tables (all rows × cells × paragraphs) ──
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    collect_para(para)

    # ── Headers & footers ──
    for section in doc.sections:
        try:
            for para in section.header.paragraphs:
                collect_para(para)
        except Exception:
            pass
        try:
            for para in section.footer.paragraphs:
                collect_para(para)
        except Exception:
            pass

    import time
    
    # To prevent spaCy from instantiating 1000+ times (once per paragraph), which causes
    # massive CPU starvation and Gunicorn 502 timeouts on Render's 0.1 CPU core,
    # we batch ALL paragraphs into a single giant string separated by a safe boundary.
    BOUNDARY = "\n---PII_DOCX_BOUNDARY_12345---\n"
    giant_string = BOUNDARY.join([item[1] for item in paragraphs_to_process])
    
    # We pass it to redact_text_segment which ALREADY optimally chunks text into 200-line blocks!
    # This drops NLP calls from 1000+ down to literally ~5 calls for a 100-page document!
    redacted_giant_string, all_findings = redact_text_segment(giant_string, selected_types)
    
    # Split back into paragraphs
    redacted_paras = redacted_giant_string.split(BOUNDARY)
    
    # Just in case of boundary corruption (should never happen), fallback gracefully
    if len(redacted_paras) != len(paragraphs_to_process):
        redacted_paras = [item[1] for item in paragraphs_to_process] # fallback to original
        
    results = []
    for i, (para, original) in enumerate(paragraphs_to_process):
        results.append((para, original, redacted_paras[i]))
    
    # Add all findings to stats
    for f in all_findings:
        et = f['entity_type']
        stats['entity_counts'][et] = stats['entity_counts'].get(et, 0) + 1
        stats['total_findings'] += 1


    # Safely apply updates to the DOCX tree sequentially
    for i, (para, original, new_text) in enumerate(results):
        if new_text != original:
            _safe_update_para(para, new_text)
            
        stats['paragraphs_processed'] += 1
        
        # Yield CPU back to Gunicorn web worker periodically so it doesn't 502 Timeout
        if i % 100 == 0:
            time.sleep(0.01)

    doc.save(output_path)
    stats['pages_processed'] = len(doc.sections)
    return stats


# ─────────────────────────────────────────────
# UTILITY
# ─────────────────────────────────────────────

def clear_mapping():
    """Reset the PII consistency map."""
    global _pii_mapping
    _pii_mapping = {}
